"""블로그 글을 Word(.docx)로 저장 — 네이버 블로그에 복붙 최적화."""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMAGE_CSV_PATH = Path("data/image/image_list.csv")


def _load_image_links() -> dict:
    """image_list.csv에서 이미지명→링크 매핑을 로드."""
    import csv
    links = {}
    if IMAGE_CSV_PATH.exists():
        with open(IMAGE_CSV_PATH, encoding="utf-8-sig") as f:
            for row in csv.DictReader(f):
                name = row.get("이미지이름", "")
                link = row.get("링크주소", "").strip()
                if name and link:
                    links[name] = link
    return links


def _get_link_for_image(img_path: str, links: dict) -> str:
    """이미지 경로에 매칭되는 링크 URL을 찾는다."""
    filename = Path(img_path).name
    return links.get(filename, "")


def _add_image_if_exists(doc, img_path: str):
    """로컬 이미지 파일이 있으면 Word에 삽입."""
    if not img_path or img_path in (".", ""):
        return False
    if img_path.startswith("IMGUR:") or img_path.startswith("http"):
        return False

    p = Path(img_path)
    if p.exists() and p.is_file():
        try:
            doc.add_picture(str(p), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception:
            return False
    return False


def export_to_docx(blog_data: dict, output_path: str = "") -> str:
    """블로그 데이터를 Word 파일로 저장한다."""
    doc = Document()
    links = _load_image_links()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    body = blog_data.get("body", "")
    title = blog_data.get("title", "")

    # ── 제목 ──
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(16)

    lines = body.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # ── 이미지 태그 처리 ──
        img_match = re.match(r"\[(썸네일 이미지|서비스 이미지|CTA 이미지):\s*(.+?)\]", stripped)
        if img_match:
            img_type = img_match.group(1)
            img_path = img_match.group(2).strip()

            # 실제 이미지 삽입
            inserted = False
            if img_path.startswith("IMGUR:"):
                # 썸네일 → 로컬 파일에서 가져오기
                thumb_dir = Path("data/generated/thumbnails")
                if thumb_dir.exists():
                    thumbs = sorted(thumb_dir.glob("*.png"), key=lambda f: f.stat().st_mtime, reverse=True)
                    if thumbs:
                        inserted = _add_image_if_exists(doc, str(thumbs[0]))
            elif not img_path.startswith("http"):
                inserted = _add_image_if_exists(doc, img_path)

            # 이미지에 매칭되는 링크가 있으면 표시
            link_url = _get_link_for_image(img_path, links)
            if link_url:
                p = doc.add_paragraph()
                run = p.add_run(f"↑ 이미지 링크: {link_url}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(180, 180, 180)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")
            i += 1
            continue

        # ── 구분선 ──
        if stripped == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # ── H2 소제목 ──
        if stripped.startswith("## "):
            text = stripped[3:].strip("* ")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # ── H3 소제목 ──
        if stripped.startswith("### "):
            text = stripped[4:].strip("* ")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # ── H1 ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip("* ")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(24)
            i += 1
            continue

        # ── 해시태그 ──
        if stripped.startswith("#") and stripped.count("#") >= 3:
            hashtags = stripped.replace("#", " #").strip()
            p = doc.add_paragraph()
            run = p.add_run(hashtags)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(59, 148, 250)
            i += 1
            continue

        # ── 일반 텍스트 ──
        p = doc.add_paragraph()
        _add_rich_text(p, stripped)
        i += 1

    # ── 저장 ──
    if not output_path:
        from datetime import datetime
        today = datetime.now().strftime("%Y%m%d")
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:30]
        out_dir = Path("data/generated")
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = str(out_dir / f"{today}_{safe_title}.docx")

    doc.save(output_path)
    print(f"  [Word] 저장: {output_path}")
    return output_path


def _add_rich_text(paragraph, text: str):
    """**bold** 마크다운을 Word 볼드로 변환."""
    segments = re.split(r"\*\*(.*?)\*\*", text)
    for idx, seg in enumerate(segments):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        run.font.size = Pt(11)
        if idx % 2 == 1:
            run.bold = True
