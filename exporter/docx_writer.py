"""블로그 포스트 → Word(.docx) 저장
BlogPilot(florence449) 구조 참고 + 마크클라우드 파이프라인에 맞게 수정
"""
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

OUTPUT_DIR = Path("output")


def _set_font(run, size_pt: int = 11, bold: bool = False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run.bold = bold


def _add_heading(doc, text: str, level: int = 2):
    sizes = {1: 18, 2: 15, 3: 13}
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    _set_font(run, size_pt=sizes.get(level, 13), bold=True)
    return p


def _add_para(doc, text: str, italic: bool = False, size_pt: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt)
    run.italic = italic
    return p


def export_docx(post: dict) -> Path | None:
    """post dict → output/{날짜}_{서비스}_{제목}.docx 저장
    성공 시 Path 반환, 실패 시 None 반환
    """
    if not _DOCX_AVAILABLE:
        print("  [docx] python-docx 미설치 — 'pip install python-docx'")
        return None

    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        doc = Document()

        # ── 제목
        _add_heading(doc, post.get("title", "(제목 없음)"), level=1)

        # ── 메타설명
        meta = post.get("meta_description", "")
        if meta:
            _add_para(doc, f"[메타설명] {meta}", italic=True, size_pt=10)

        doc.add_paragraph()  # 빈 줄

        # ── 본문 (구조형: intro + body sections + conclusion + cta)
        if isinstance(post.get("body"), list):
            intro = post.get("intro", "")
            if intro:
                _add_para(doc, intro)
                doc.add_paragraph()

            for section in post.get("body", []):
                if isinstance(section, dict):
                    heading = section.get("heading", "")
                    content = section.get("content", "")
                    if heading:
                        _add_heading(doc, heading, level=2)
                    if content:
                        _add_para(doc, content)
                        doc.add_paragraph()

            conclusion = post.get("conclusion", "")
            if conclusion:
                _add_heading(doc, "마무리", level=2)
                _add_para(doc, conclusion)
                doc.add_paragraph()

            cta = post.get("cta", "")
            if cta:
                p = doc.add_paragraph()
                run = p.add_run(f"▶ {cta}")
                _set_font(run, bold=True)

        else:
            # 구형 flat body 폴백
            body = post.get("body", "")
            if body:
                _add_para(doc, body)

        # ── 해시태그
        tags = post.get("hashtags", [])
        if tags:
            doc.add_paragraph()
            tag_str = " ".join(t if t.startswith("#") else f"#{t}" for t in tags)
            p = doc.add_paragraph()
            run = p.add_run(tag_str)
            _set_font(run, size_pt=10)
            run.italic = True

        # ── 파일명 생성: YYYYMMDD_서비스_제목(30자).docx
        date_str   = datetime.now().strftime("%Y%m%d")
        svc        = post.get("service_key", "blog")
        title_safe = re.sub(r'[\\/:*?"<>|\u2014\u2013]', "_", post.get("title", "post"))[:30]
        filepath   = OUTPUT_DIR / f"{date_str}_{svc}_{title_safe}.docx"

        doc.save(filepath)
        print(f"  [docx] 저장: {filepath}")
        return filepath

    except Exception as e:
        print(f"  [docx] 저장 실패: {e}")
        return None
